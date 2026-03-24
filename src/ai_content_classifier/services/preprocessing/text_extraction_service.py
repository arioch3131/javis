# services/preprocessing/text_extraction_service.py
"""
Fast text extraction service for LLMs.

This service specializes in rapid text extraction from various formats
for sending to LLM models. It prioritizes speed and simplicity:
- Basic extraction without metadata
- Configurable length limitation
- Intelligent fallbacks
- Caching to avoid re-extraction

It does NOT perform full metadata extraction (delegated to MetadataService).
"""

import hashlib
import os
import time
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional

from ai_content_classifier.core.logger import LoggableMixin
from ai_content_classifier.services.shared.cache_runtime import get_cache_runtime


@dataclass
class TextExtractionConfig:
    """Configuration pour l'extraction de texte."""

    # Content limits
    max_text_length: int = 50000  # Max characters for LLM
    max_pages_pdf: int = 50  # Max pages to process
    preview_length: int = 1000  # Preview length for debug

    # Encodings to test
    encoding_fallbacks: List[str] = field(
        default_factory=lambda: ["utf-8", "latin-1", "cp1252", "iso-8859-1"]
    )

    # Cache and performance
    cache_size: int = 100
    extraction_timeout: int = 30  # Max seconds per file

    # Processing options
    clean_whitespace: bool = True
    preserve_line_breaks: bool = True
    strip_binary_artifacts: bool = True


@dataclass
class TextExtractionResult:
    """Result of text extraction."""

    success: bool
    text: str = ""

    # Extraction information
    extraction_method: str = ""
    original_length: int = 0
    final_length: int = 0
    truncated: bool = False
    processing_time: float = 0.0

    # Basic metadata
    file_format: str = ""
    pages_processed: int = 0
    encoding_used: str = ""

    # Diagnostics
    warnings: List[str] = field(default_factory=list)
    error_message: Optional[str] = None

    @property
    def is_empty(self) -> bool:
        """Checks if the extracted text is empty."""
        return not self.text.strip()

    @property
    def truncation_ratio(self) -> float:
        """Truncation ratio (0.0 = not truncated, 1.0 = completely truncated)."""
        if self.original_length == 0:
            return 0.0
        return 1.0 - (self.final_length / self.original_length)


class TextExtractionService(LoggableMixin):
    """
    Text extraction service optimized for LLM workflows.

    This service extracts text quickly from multiple formats
    to send it to language models. It prioritizes:
    - Speed over completeness
    - Simplicity over metadata richness
    - La robustesse avec des fallbacks
    """

    def __init__(self, config: Optional[TextExtractionConfig] = None):
        """
        Initialise le service d'extraction de texte.

        Args:
            config: Configuration optionnelle
        """
        super().__init__()
        self.__init_logger__()

        self.config = config or TextExtractionConfig()

        self._cache_runtime = get_cache_runtime()
        self._extraction_cache = self._cache_runtime.memory_cache(
            "preprocessing:text_extraction",
            default_ttl=600,
        )

        # Check available dependencies
        self._available_extractors = self._check_available_extractors()

        # Statistiques
        self._stats = {
            "extractions_performed": 0,
            "cache_hits": 0,
            "total_processing_time": 0.0,
            "successful_extractions": 0,
            "failed_extractions": 0,
            "formats_processed": {},
        }

        self.logger.info(
            f"TextExtractionService initialized - Available extractors: {list(self._available_extractors.keys())}"
        )

    def _check_available_extractors(self) -> Dict[str, bool]:
        """Check which extraction libraries are available."""

        available = {}

        # pypdf pour PDF
        try:
            import pypdf  # noqa: F401

            available["pypdf"] = True
        except ImportError:
            available["pypdf"] = False

        # pdfminer pour PDF (alternative)
        try:
            from pdfminer.high_level import extract_text  # noqa: F401

            available["pdfminer"] = True
        except ImportError:
            available["pdfminer"] = False

        # python-docx pour DOCX
        try:
            import docx  # noqa: F401

            available["docx"] = True
        except ImportError:
            available["docx"] = False

        # striprtf pour RTF
        try:
            import striprtf.striprtf  # noqa: F401

            available["striprtf"] = True
        except ImportError:
            available["striprtf"] = False

        # textract (solution universelle mais plus lente)
        try:
            import textract  # noqa: F401

            available["textract"] = True
        except ImportError:
            available["textract"] = False

        return available

    def extract_text_for_llm(
        self, file_path: str, max_length: Optional[int] = None
    ) -> TextExtractionResult:
        """
        Extrait le texte d'un file pour envoi vers un LLM.

        Args:
            file_path: Chemin vers le file
            max_length: Longueur max override (utilise config si None)

        Returns:
            TextExtractionResult avec texte extrait
        """
        start_time = time.time()
        max_length = max_length or self.config.max_text_length

        # Check le cache
        cache_key = self._generate_cache_key(file_path, max_length)
        cached_result = self._extraction_cache.get(cache_key, default=None)
        if cached_result is not None:
            self._stats["cache_hits"] += 1
            self.logger.debug(f"Cache hit pour {file_path}")
            return cached_result

        try:
            # Determine file format
            file_format = self._detect_file_format(file_path)

            # Extraire selon le format
            result = self._extract_by_format(
                file_path, file_format, max_length, start_time
            )

            # Mise en cache si success
            if result.success:
                self._cache_result(cache_key, result)

            # Update statistics
            self._update_stats(result, file_format)

            return result

        except Exception as e:
            processing_time = time.time() - start_time
            error_msg = f"Erreur extraction texte {file_path}: {str(e)}"

            result = TextExtractionResult(
                success=False,
                error_message=error_msg,
                processing_time=processing_time,
                file_format=self._detect_file_format(file_path),
            )

            self.logger.error(error_msg, exc_info=True)
            self._stats["failed_extractions"] += 1
            return result

    def _detect_file_format(self, file_path: str) -> str:
        """Detect file format based on extension."""

        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        format_map = {
            ".txt": "text",
            ".md": "markdown",
            ".csv": "csv",
            ".pdf": "pdf",
            ".doc": "doc",
            ".docx": "docx",
            ".rtf": "rtf",
            ".odt": "odt",
        }

        return format_map.get(ext, "unknown")

    def _extract_by_format(
        self, file_path: str, file_format: str, max_length: int, start_time: float
    ) -> TextExtractionResult:
        """Extrait le texte selon le format de file."""

        # Extraction methods by format
        extraction_methods = {
            "text": self._extract_plain_text,
            "markdown": self._extract_plain_text,
            "csv": self._extract_plain_text,
            "pdf": self._extract_pdf_text,
            "docx": self._extract_docx_text,
            "rtf": self._extract_rtf_text,
            "doc": self._extract_legacy_doc,
            "odt": self._extract_odt_text,
        }

        # Try extraction with the appropriate method
        extraction_method = extraction_methods.get(file_format, self._extract_fallback)

        try:
            text, method_name, pages_processed, encoding_used, warnings = (
                extraction_method(file_path)
            )

            if text:
                # Post-traitement du texte
                processed_text = self._post_process_text(text, max_length)

                processing_time = time.time() - start_time

                result = TextExtractionResult(
                    success=True,
                    text=processed_text,
                    extraction_method=method_name,
                    original_length=len(text),
                    final_length=len(processed_text),
                    truncated=len(processed_text) < len(text),
                    processing_time=processing_time,
                    file_format=file_format,
                    pages_processed=pages_processed,
                    encoding_used=encoding_used,
                    warnings=warnings,
                )

                self.logger.debug(
                    f"Texte extrait: {len(text)} -> {len(processed_text)} chars, "
                    f"method: {method_name}, temps: {processing_time:.2f}s"
                )

                return result
            else:
                return TextExtractionResult(
                    success=False,
                    error_message="No texte extrait",
                    processing_time=time.time() - start_time,
                    file_format=file_format,
                    extraction_method=method_name,
                )

        except Exception as e:
            return TextExtractionResult(
                success=False,
                error_message=f"Erreur {extraction_method.__name__}: {str(e)}",
                processing_time=time.time() - start_time,
                file_format=file_format,
            )

    def _extract_plain_text(self, file_path: str) -> tuple:
        """Extrait le texte depuis un file texte plain."""

        for encoding in self.config.encoding_fallbacks:
            try:
                with open(file_path, "r", encoding=encoding, errors="ignore") as f:
                    text = f.read()

                return text, "plain_text", 1, encoding, []

            except UnicodeDecodeError:
                continue
            except Exception as e:
                self.logger.warning(f"Erreur lecture file texte avec {encoding}: {e}")
                continue

        raise Exception("Unable to read file with supported encodings")

    def _extract_pdf_text(self, file_path: str) -> tuple:
        """Extrait le texte depuis un PDF."""

        # Essayer pypdf en premier (plus rapide)
        if self._available_extractors.get("pypdf"):
            try:
                text, pages = self._extract_pdf_pypdf(file_path)
                if text.strip():
                    return text, "pypdf", pages, "utf-8", []
            except Exception as e:
                self.logger.debug(f"pypdf failure: {e}")

        # Fallback vers pdfminer (plus robuste)
        if self._available_extractors.get("pdfminer"):
            try:
                text = self._extract_pdf_pdfminer(file_path)
                if text.strip():
                    return (
                        text,
                        "pdfminer",
                        -1,
                        "utf-8",
                        ["Pages not counted with pdfminer"],
                    )
            except Exception as e:
                self.logger.debug(f"pdfminer failure: {e}")

        # Fallback textract
        if self._available_extractors.get("textract"):
            try:
                text = self._extract_with_textract(file_path)
                if text.strip():
                    return (
                        text,
                        "textract",
                        -1,
                        "utf-8",
                        ["Extraction textract (lente)"],
                    )
            except Exception as e:
                self.logger.debug(f"textract failure: {e}")

        raise Exception("Noe method d'extraction PDF disponible")

    def _extract_pdf_pypdf(self, file_path: str) -> tuple:
        """Extraction PDF avec pypdf."""
        import pypdf

        text_parts = []

        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            total_pages = len(reader.pages)
            pages_to_read = min(total_pages, self.config.max_pages_pdf)

            for page_num in range(pages_to_read):
                try:
                    page = reader.pages[page_num]
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_parts.append(page_text)
                except Exception as e:
                    self.logger.debug(f"Erreur page {page_num}: {e}")
                    continue

        return "\n\n".join(text_parts), pages_to_read

    def _extract_pdf_pdfminer(self, file_path: str) -> str:
        """Extraction PDF avec pdfminer."""
        from pdfminer.high_level import extract_text

        return extract_text(file_path)

    def _extract_docx_text(self, file_path: str) -> tuple:
        """Extrait le texte depuis un file DOCX."""

        if not self._available_extractors.get("docx"):
            return (
                self._extract_with_textract(file_path),
                "textract_fallback",
                -1,
                "utf-8",
                ["Fallback textract pour DOCX"],
            )

        try:
            import docx

            doc = docx.Document(file_path)
            text_parts = []

            # Extraire les paragraphes
            for para in doc.paragraphs:
                if para.text:
                    text_parts.append(para.text)

            # Extraire les tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text:
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return (
                "\n\n".join(text_parts),
                "python_docx",
                len(doc.paragraphs),
                "utf-8",
                [],
            )

        except Exception as e:
            if self._available_extractors.get("textract"):
                text = self._extract_with_textract(file_path)
                return (
                    text,
                    "textract_fallback",
                    -1,
                    "utf-8",
                    [f"Fallback textract: {str(e)}"],
                )
            else:
                raise

    def _extract_rtf_text(self, file_path: str) -> tuple:
        """Extrait le texte depuis un file RTF."""

        if self._available_extractors.get("striprtf"):
            try:
                import striprtf.striprtf

                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    rtf_content = f.read()

                text = striprtf.striprtf.rtf_to_text(rtf_content)
                return text, "striprtf", 1, "utf-8", []

            except Exception as e:
                self.logger.debug(f"striprtf failure: {e}")

        # Fallback vers extraction texte brut (basique)
        try:
            text, _, pages, encoding, warnings = self._extract_plain_text(file_path)
            # Nettoyage basique des codes RTF
            text = self._clean_rtf_artifacts(text)
            warnings.append("Extraction RTF basique - formatage perdu")
            return text, "rtf_basic", pages, encoding, warnings

        except Exception:
            if self._available_extractors.get("textract"):
                text = self._extract_with_textract(file_path)
                return (
                    text,
                    "textract_fallback",
                    -1,
                    "utf-8",
                    ["Fallback textract pour RTF"],
                )
            else:
                raise

    def _extract_legacy_doc(self, file_path: str) -> tuple:
        """Extrait le texte depuis un file DOC legacy."""

        if self._available_extractors.get("textract"):
            text = self._extract_with_textract(file_path)
            return text, "textract", -1, "utf-8", ["Extraction DOC legacy via textract"]
        else:
            raise Exception("Legacy DOC extraction requires textract")

    def _extract_odt_text(self, file_path: str) -> tuple:
        """Extrait le texte depuis un file ODT."""

        if self._available_extractors.get("textract"):
            text = self._extract_with_textract(file_path)
            return text, "textract", -1, "utf-8", ["Extraction ODT via textract"]
        else:
            raise Exception("ODT extraction requires textract")

    def _extract_with_textract(self, file_path: str) -> str:
        """Extraction avec textract (solution universelle mais lente)."""
        import textract

        text_bytes = textract.process(file_path)
        return text_bytes.decode("utf-8", errors="ignore")

    def _extract_fallback(self, file_path: str) -> tuple:
        """Last-resort extraction method."""

        # Essayer textract si disponible
        if self._available_extractors.get("textract"):
            try:
                text = self._extract_with_textract(file_path)
                return (
                    text,
                    "textract_fallback",
                    -1,
                    "utf-8",
                    ["Extraction fallback textract"],
                )
            except Exception as e:
                self.logger.debug(f"textract fallback failure: {e}")

        # Extraction brute comme texte
        try:
            text, _, pages, encoding, warnings = self._extract_plain_text(file_path)
            text = self._clean_binary_artifacts(text)
            warnings.append("Raw extraction - result not guaranteed")
            return text, "raw_fallback", pages, encoding, warnings
        except Exception as e:
            raise Exception(f"All extraction methods failed: {str(e)}")

    def _post_process_text(self, text: str, max_length: int) -> str:
        """Post-traite le texte extrait."""

        if self.config.clean_whitespace:
            # Nettoyage des espaces
            import re

            if self.config.preserve_line_breaks:
                # Keep line breaks but clean spaces
                text = re.sub(r" +", " ", text)  # Espaces multiples -> simple
                text = re.sub(
                    r"\n\s*\n", "\n\n", text
                )  # Lignes vides multiples -> double
            else:
                # Nettoyage agressif
                text = re.sub(r"\s+", " ", text)

        # Remove binary artifacts if requested
        if self.config.strip_binary_artifacts:
            text = self._clean_binary_artifacts(text)

        # Truncate if needed
        if len(text) > max_length:
            text = text[:max_length]
            # Try cutting at sentence or line boundary
            last_period = text.rfind(".")
            last_newline = text.rfind("\n")
            cut_point = max(last_period, last_newline)

            if cut_point > max_length * 0.8:  # Si on peut couper proprement
                text = text[: cut_point + 1]

        return text.strip()

    def _clean_rtf_artifacts(self, text: str) -> str:
        """Nettoie les artifacts RTF basiques."""
        import re

        # Supprimer les commandes RTF communes
        text = re.sub(r"\\[a-z]+\d*", "", text)
        text = re.sub(r"[{}]", "", text)
        text = re.sub(r"\\\*.*?;", "", text)

        return text

    def _clean_binary_artifacts(self, text: str) -> str:
        """Nettoie les artifacts binaires."""

        # Keep only printable characters and spaces
        cleaned = "".join(char for char in text if char.isprintable() or char.isspace())

        return cleaned

    def _generate_cache_key(self, file_path: str, max_length: int) -> str:
        """Generate cache key."""

        try:
            mtime = os.path.getmtime(file_path)
        except Exception:
            mtime = time.time()

        key_data = f"{file_path}_{max_length}_{mtime}"
        return hashlib.md5(key_data.encode()).hexdigest()

    def _cache_result(self, cache_key: str, result: TextExtractionResult):
        """Cache the result."""

        self._extraction_cache.set(cache_key, result)

    def _update_stats(self, result: TextExtractionResult, file_format: str):
        """Update les statistiques."""

        self._stats["extractions_performed"] += 1
        self._stats["total_processing_time"] += result.processing_time

        if result.success:
            self._stats["successful_extractions"] += 1
        else:
            self._stats["failed_extractions"] += 1

        # Statistiques par format
        if file_format not in self._stats["formats_processed"]:
            self._stats["formats_processed"][file_format] = 0
        self._stats["formats_processed"][file_format] += 1

    # === MÉTHODES UTILITAIRES PUBLIQUES ===

    def get_supported_formats(self) -> List[str]:
        """Return list of supported formats."""

        base_formats = ["txt", "md", "csv"]

        if self._available_extractors.get("pypdf") or self._available_extractors.get(
            "pdfminer"
        ):
            base_formats.append("pdf")

        if self._available_extractors.get("docx"):
            base_formats.append("docx")

        if self._available_extractors.get("striprtf"):
            base_formats.append("rtf")

        if self._available_extractors.get("textract"):
            base_formats.extend(
                ["doc", "odt", "docx", "rtf", "pdf"]
            )  # textract supporte beaucoup

        return list(set(base_formats))

    def clear_cache(self):
        """Vide le cache d'extraction."""
        self._extraction_cache.clear()
        self.logger.info("Extraction cache cleared")

    def get_stats(self) -> Dict[str, Any]:
        """Return les statistiques."""

        avg_processing_time = self._stats["total_processing_time"] / max(
            1, self._stats["extractions_performed"]
        )

        success_rate = self._stats["successful_extractions"] / max(
            1, self._stats["extractions_performed"]
        )

        return {
            "extractions_performed": self._stats["extractions_performed"],
            "cache_hits": self._stats["cache_hits"],
            "success_rate": success_rate,
            "avg_processing_time": avg_processing_time,
            "supported_formats": self.get_supported_formats(),
            "available_extractors": self._available_extractors,
            "formats_processed": self._stats["formats_processed"],
            "cache_size": self._extraction_cache.size(),
        }
