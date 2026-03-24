"""
Hachoir metadata extractor module.

This module provides metadata extraction using the Hachoir library.
"""

import importlib
import os
from pathlib import Path
from typing import Any, Dict, List, Optional

from ai_content_classifier.services.metadata.extractors.base_extractor import (
    BaseMetadataExtractor,
)


class HachoirExtractor(BaseMetadataExtractor):
    """Generic metadata extractor using Hachoir."""

    # Map of common Hachoir metadata groups and fields of interest
    IMPORTANT_METADATA = {
        "Common": ["width", "height", "duration", "bit_rate", "comment", "mime_type"],
        "Audio": ["sample_rate", "bits_per_sample", "channel_count"],
        "Video": ["frame_rate", "compression"],
    }

    # File extension priority groups
    EXTENSION_GROUPS = {
        "documents": [
            ".pdf",
            ".doc",
            ".docx",
            ".xls",
            ".xlsx",
            ".ppt",
            ".pptx",
            ".odt",
            ".ods",
            ".odp",
            ".rtf",
        ],
        "archives": [".zip", ".rar", ".7z", ".tar", ".gz"],
        "media": [".mp3", ".mp4", ".avi", ".mov", ".mkv", ".flac", ".wav"],
    }

    def __init__(self, additional_extensions: Optional[List[str]] = None):
        """
        Initialize the Hachoir extractor.

        Args:
            additional_extensions: Optional list of additional file extensions to handle
        """
        super().__init__()

        # Flatten the extension groups into a single priority list
        self.priority_extensions = []
        for group in self.EXTENSION_GROUPS.values():
            self.priority_extensions.extend(group)

        # Add any additional extensions provided
        if additional_extensions:
            self.priority_extensions.extend(additional_extensions)

        # Check if required libraries are available
        self.hachoir_available = self._check_dependency(
            "hachoir.parser"
        ) and self._check_dependency("hachoir.metadata")
        self.pypdf_available = self._check_dependency("pypdf")
        self.docx_available = self._check_dependency("docx")

        if not self.hachoir_available:
            self.logger.warning(
                "Hachoir library not available. Some functionality will be limited."
            )

    def _check_dependency(self, module_name: str) -> bool:
        """
        Check if a Python module is available.

        Args:
            module_name: Name of the module to check

        Returns:
            True if the module is available, False otherwise
        """
        try:
            importlib.import_module(module_name)
            return True
        except ImportError:
            return False

    def can_handle(self, file_path: str) -> bool:
        """
        Hachoir can try to process almost any file,
        but we prioritize certain extensions.

        Args:
            file_path: Path to the file

        Returns:
            True if Hachoir is the recommended extractor, False otherwise
        """
        if not self.hachoir_available:
            return False

        ext = Path(file_path).suffix.lower()

        # Check if the file exists and is readable
        if not os.path.isfile(file_path) or not os.access(file_path, os.R_OK):
            return False

        # We prioritize certain extensions
        return ext in self.priority_extensions

    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from the file using Hachoir.

        Args:
            file_path: Path to the file

        Returns:
            Dictionary of metadata
        """
        # Get basic metadata
        metadata = self.get_basic_metadata(file_path)

        if not self.hachoir_available:
            metadata["error"] = "Hachoir library is not installed"
            return metadata

        try:
            # Extract Hachoir metadata
            hachoir_metadata = self._extract_hachoir_metadata(file_path)
            if hachoir_metadata:
                metadata.update(hachoir_metadata)

            # Extract format-specific metadata if available
            mime_type = metadata.get("mime_type", "")

            # PDF-specific extraction
            if "application/pdf" in mime_type and self.pypdf_available:
                pdf_metadata = self._extract_pdf_metadata(file_path)
                if pdf_metadata:
                    metadata["pdf_info"] = pdf_metadata

            # DOCX-specific extraction
            if (
                "officedocument.wordprocessingml.document" in mime_type
                and self.docx_available
            ):
                docx_metadata = self._extract_docx_metadata(file_path)
                if docx_metadata:
                    metadata["docx_info"] = docx_metadata

        except Exception as e:
            self.logger.error(f"Error extracting metadata: {str(e)}")
            metadata["error"] = f"Error extracting metadata: {str(e)}"

        return metadata

    def _extract_hachoir_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata using the Hachoir library.

        Args:
            file_path: Path to the file

        Returns:
            Dictionary of extracted metadata
        """
        from hachoir.metadata import extractMetadata
        from hachoir.parser import createParser

        metadata = {}

        # Create the parser
        parser = createParser(file_path)
        if not parser:
            self.logger.warning(f"Unable to parse file with Hachoir: {file_path}")
            return metadata

        # Extract metadata
        hachoir_metadata = extractMetadata(parser)
        if not hachoir_metadata:
            self.logger.warning(f"No metadata extracted with Hachoir: {file_path}")
            return metadata

        # Process metadata by group
        metadata["hachoir"] = {}
        extracted_dates = {}

        for meta_group in hachoir_metadata:
            group_name = meta_group.header

            if group_name not in metadata["hachoir"]:
                metadata["hachoir"][group_name] = {}

            for meta_item in meta_group:
                self._process_metadata_item(
                    meta_item, metadata, group_name, extracted_dates
                )

        # Move important metadata to the top level for easier access
        self._promote_important_metadata(metadata)

        # Set creation date and last modified from extracted dates
        if "creation_date" in extracted_dates:
            metadata["creation_date"] = extracted_dates["creation_date"]
        if "last_modified" in extracted_dates:
            metadata["last_modified"] = extracted_dates["last_modified"]

        return metadata

    def _process_metadata_item(self, meta_item, metadata, group_name, extracted_dates):
        """
        Process a single metadata item from Hachoir.

        Args:
            meta_item: Hachoir metadata item
            metadata: Metadata dictionary to update
            group_name: Name of the metadata group
            extracted_dates: Dictionary of extracted date fields
        """
        key = meta_item.key

        if meta_item.values:
            # Extract and convert the value
            if len(meta_item.values) == 1:
                value = meta_item.values[0].value

                # Handle date values
                if (
                    hasattr(value, "year")
                    and hasattr(value, "month")
                    and hasattr(value, "day")
                ):
                    # It's probably a date
                    try:
                        if group_name == "Common" and key == "creation_date":
                            extracted_dates["creation_date"] = value
                        elif group_name == "Common" and key == "last_modification":
                            extracted_dates["last_modified"] = value
                    except Exception as e:
                        self.logger.debug(f"Error processing date: {str(e)}")

                # Store in the hachoir group
                metadata["hachoir"][group_name][key] = value
            else:
                # Multiple values
                values = [val.value for val in meta_item.values]
                metadata["hachoir"][group_name][key] = values

    def _promote_important_metadata(self, metadata):
        """
        Move important metadata to the top level for easier access.

        Args:
            metadata: Metadata dictionary to update
        """
        for group_name, fields in self.IMPORTANT_METADATA.items():
            if group_name in metadata["hachoir"]:
                for key in fields:
                    if key in metadata["hachoir"][group_name]:
                        metadata[key] = metadata["hachoir"][group_name][key]

    def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract additional metadata specific to PDF files.

        Args:
            file_path: Path to the PDF file

        Returns:
            Dictionary of PDF-specific metadata
        """
        try:
            import pypdf

            pdf_info = {}

            with open(file_path, "rb") as f:
                pdf = pypdf.PdfReader(f)

                # Add information about content
                pdf_info["pages_count"] = len(pdf.pages)

                # Extract text from first page for preview (limited to 1000 chars)
                try:
                    first_page_text = pdf.pages[0].extract_text()
                    if first_page_text:
                        pdf_info["first_page_preview"] = first_page_text[:1000]
                except Exception as e:
                    self.logger.debug(f"Could not extract text from PDF: {str(e)}")

                # Extract document metadata
                if pdf.metadata:
                    doc_info = {}
                    for key, value in pdf.metadata.items():
                        if value and isinstance(value, (str, int, float, bool)):
                            clean_key = key.replace("/", "")
                            doc_info[clean_key] = value

                    if doc_info:
                        pdf_info["document_info"] = doc_info

                return pdf_info

        except Exception as e:
            self.logger.debug(f"Error extracting PDF metadata: {str(e)}")
            return {}

    def _extract_docx_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract additional metadata specific to DOCX files.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Dictionary of DOCX-specific metadata
        """
        try:
            import docx

            docx_info = {}

            doc = docx.Document(file_path)

            # Add information about content structure
            docx_info["document_structure"] = {
                "paragraphs_count": len(doc.paragraphs),
                "tables_count": len(doc.tables),
                "sections_count": len(doc.sections),
            }

            # Add text statistics if there's content
            if doc.paragraphs:
                words = []
                characters = 0

                for p in doc.paragraphs:
                    if p.text:
                        paragraph_words = p.text.split()
                        words.extend(paragraph_words)
                        characters += len(p.text)

                docx_info["text_statistics"] = {
                    "word_count": len(words),
                    "character_count": characters,
                    "paragraph_count": len(doc.paragraphs),
                }

                # Add text preview (first 1000 chars)
                text_preview = "\n".join(p.text for p in doc.paragraphs[:5] if p.text)
                if text_preview:
                    docx_info["text_preview"] = text_preview[:1000]

            # Extract core properties if available
            try:
                core_props = doc.core_properties
                if core_props:
                    props = {}
                    for attr in dir(core_props):
                        if not attr.startswith("_") and not callable(
                            getattr(core_props, attr)
                        ):
                            value = getattr(core_props, attr)
                            if value is not None:
                                props[attr] = value

                    if props:
                        docx_info["core_properties"] = props
            except Exception as e:
                self.logger.debug(f"Could not extract core properties: {str(e)}")

            return docx_info

        except Exception as e:
            self.logger.debug(f"Error extracting DOCX metadata: {str(e)}")
            return {}
